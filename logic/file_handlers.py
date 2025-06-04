import os
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog, messagebox
import docx
from PyPDF2 import PdfReader
import threading
from logic.openai_analysis import analizar_texto

def setup_menu(app):
    menubar = app.root.config(menu=app.root)
    menubar = app.root.config(menu=app.root)
    menubar = app.root.config(menu=app.root)

    menubar = app.root.config(menu=app.root)
    menubar = app.root.config(menu=app.root)

    menubar = app.root.config(menu=app.root)


def paste_text(app):
    try:
        text = app.root.clipboard_get()
        app.text_widget.delete("1.0", "end")
        app.text_widget.insert("end", text)
        update_stats(app)
        app.current_file = None
        app.status.config(text="📋 Texto pegado del portapapeles.")
        # Lanzar análisis con Ollama
        lanzar_analisis_con_popup(app)
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo pegar el texto: {str(e)}")

def load_file(app):
    ruta = filedialog.askopenfilename(filetypes=[
        ("Archivos de texto", "*.txt"),
        ("Word", "*.docx"),
        ("PDF", "*.pdf"),
        ("Todos los archivos", "*.*")
    ])
    if not ruta:
        return
    try:
        with open(ruta, 'rb') as f:
            content = ""
            if ruta.endswith(".txt"):
                content = f.read().decode('utf-8')
            elif ruta.endswith(".docx"):
                content = "\n".join(p.text for p in docx.Document(ruta).paragraphs)
            elif ruta.endswith(".pdf"):
                reader = PdfReader(ruta)
                content = "\n".join(page.extract_text() or "" for page in reader.pages)
            else:
                content = "⚠️ Formato no compatible o bibliotecas faltantes."

        app.text_widget.delete("1.0", "end")
        app.text_widget.insert("end", content)
        update_stats(app)
        app.current_file = ruta
        app.status.config(text=f"✅ Documento cargado: {os.path.basename(ruta)}")
        # Lanzar análisis con Ollama
        lanzar_analisis_con_popup(app)
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo abrir el archivo:\n{str(e)}")

def cargar_texto_comparador(widget_destino):
    ruta = filedialog.askopenfilename(filetypes=[
        ("Archivos de texto", "*.txt"),
        ("Word", "*.docx"),
        ("PDF", "*.pdf"),
        ("Todos los archivos", "*.*")
    ])
    if not ruta:
        return
    try:
        with open(ruta, 'rb') as f:
            content = ""
            if ruta.endswith(".txt"):
                content = f.read().decode('utf-8')
            elif ruta.endswith(".docx"):
                content = "\n".join(p.text for p in docx.Document(ruta).paragraphs)
            elif ruta.endswith(".pdf"):
                reader = PdfReader(ruta)
                content = "\n".join(page.extract_text() or "" for page in reader.pages)
            else:
                content = "⚠️ Formato no compatible o bibliotecas faltantes."

        widget_destino.delete("1.0", "end")
        widget_destino.insert("1.0", content)
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo abrir el archivo:\n{str(e)}")

def save_file(app):
    path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Archivo de texto", "*.txt"), ("Todos los archivos", "*.*")])
    if path:
        try:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(app.text_widget.get("1.0", "end"))
            app.current_file = path
            app.status.config(text=f"Documento guardado en: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar el archivo:\n{str(e)}")

def clear_text(app):
    app.text_widget.delete("1.0", "end")
    update_stats(app)
    app.status.config(text="🧹 Texto limpiado.")

def update_stats(app):
    import re
    text = app.text_widget.get("1.0", "end")
    words = len(re.findall(r'\b\w+\b', text))
    chars = len(text.replace("\n", ""))
    lines = text.count("\n") + (0 if text.endswith("\n") else 1)
    sentences = len(re.findall(r'[.!?]+', text))

    app.word_label.config(text=f"Palabras: {words}")
    app.char_label.config(text=f"Caracteres: {chars}")
    app.line_label.config(text=f"Líneas: {lines}")
    app.sentence_label.config(text=f"Oraciones: {sentences}")

def mostrar_analisis(app, texto):
    for widget in app.analysis_tab.winfo_children():
        widget.destroy()

    contenedor = ttk.Frame(app.analysis_tab)
    contenedor.pack(fill=tk.BOTH, expand=True)

    caja = tk.Text(contenedor, wrap=tk.WORD, font=("Segoe UI", 11), bg="white")
    
    # Inicializa partes
    plagio_line = ""
    ia_line = ""
    fuente_line = ""

    for line in texto.splitlines():
        if "Plagio:" in line:
            plagio_line = line.strip()
        elif "Generado por IA:" in line:
            ia_line = line.strip()
        elif "Fuente probable:" in line:
            fuente_line = line.strip()

    # Inserta resultados limpios
    if plagio_line:
        caja.insert("end", f"{plagio_line}\n")
    if ia_line:
        caja.insert("end", f"{ia_line}\n")

    # Comentario según IA
    try:
        ia_score = int(ia_line.split(":")[1].strip().replace("%", ""))
        if ia_score > 80:
            comentario = "✍️ Este texto parece altamente generado por IA por su estilo uniforme y precisión sintáctica."
        elif ia_score > 50:
            comentario = "✍️ El contenido muestra indicios de generación automática combinados con rasgos humanos."
        else:
            comentario = "✍️ El estilo del texto sugiere que ha sido mayormente escrito por una persona."
        caja.insert("end", f"\n\n{comentario}")
    except:
        pass

    # Comentario de fuente
    if fuente_line:
        fuente_texto = fuente_line.split("Fuente probable:")[1].strip()
        caja.insert("end", f"\n\n🔎 Fuente probable: {fuente_texto}")

    caja.config(state=tk.DISABLED)
    caja.pack(fill=tk.BOTH, expand=True, padx=10, pady=(10, 0))

    ttk.Button(contenedor, text="Volver a analizar", command=lambda: lanzar_analisis_con_popup(app), style="Secondary.TButton").pack(pady=10)

def lanzar_analisis_con_popup(app):
    popup = tk.Toplevel(app.root)
    popup.title("Analizando...")
    popup.geometry("300x100")
    popup.resizable(False, False)
    popup.grab_set()
    ttk.Label(popup, text="Analizando con IA...\nEsto puede tardar unos segundos.", anchor="center").pack(pady=10)
    progress = ttk.Progressbar(popup, mode="indeterminate")
    progress.pack(fill=tk.X, padx=20, pady=10)
    progress.start()

    def tarea():
        try:
            texto = app.text_widget.get("1.0", "end")
            resultado = analizar_texto(texto)
        except Exception as e:
            resultado = f"Error al conectar con Ollama:\n{str(e)}"

        popup.destroy()

        if resultado.startswith("Error"):
            messagebox.showerror("Error de conexión", resultado)
        mostrar_analisis(app, resultado)

    hilo = threading.Thread(target=tarea)
    hilo.start()


